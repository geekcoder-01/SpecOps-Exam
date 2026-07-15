from pathlib import Path
from threading import Lock
from typing import Optional

import easyocr
import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from PIL import Image


SUPPORTED_IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
}

OCR_LANGUAGES = ["en"]

MIN_DIGITAL_PDF_TEXT_LENGTH = 30
PDF_RENDER_SCALE = 2.0

_ocr_reader: Optional[easyocr.Reader] = None
_ocr_reader_lock = Lock()


class TextExtractionError(Exception):
    """Raised when readable text cannot be extracted from a file."""


def get_ocr_reader() -> easyocr.Reader:
    global _ocr_reader

    if _ocr_reader is None:
        with _ocr_reader_lock:
            if _ocr_reader is None:
                try:
                    _ocr_reader = easyocr.Reader(
                        OCR_LANGUAGES,
                        gpu=False,
                    )
                except Exception as error:
                    raise TextExtractionError(
                        f"Unable to initialize EasyOCR: {error}"
                    ) from error

    return _ocr_reader


def extract_text(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise TextExtractionError(
            "Uploaded source file was not found"
        )

    extension = path.suffix.lower()

    if extension == ".pdf":
        text = extract_pdf_text(path)

    elif extension == ".docx":
        text = extract_docx_text(path)

    elif extension == ".doc":
        raise TextExtractionError(
            "Legacy DOC files are not supported directly. "
            "Please convert the file to DOCX and upload it again."
        )

    elif extension in SUPPORTED_IMAGE_EXTENSIONS:
        text = extract_image_text(path)

    else:
        raise TextExtractionError(
            f"Unsupported file format: {extension}"
        )

    cleaned_text = clean_extracted_text(text)

    if not cleaned_text:
        raise TextExtractionError(
            "No readable text could be extracted from the file"
        )

    return cleaned_text


def extract_pdf_text(path: Path) -> str:
    digital_text = extract_digital_pdf_text(path)

    if len(digital_text.strip()) >= MIN_DIGITAL_PDF_TEXT_LENGTH:
        return digital_text

    return extract_scanned_pdf_text(path)


def extract_digital_pdf_text(path: Path) -> str:
    extracted_pages = []

    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_number, page in enumerate(
                pdf.pages,
                start=1,
            ):
                try:
                    page_text = page.extract_text() or ""
                except Exception as error:
                    raise TextExtractionError(
                        f"Unable to read PDF page "
                        f"{page_number}: {error}"
                    ) from error

                if page_text.strip():
                    extracted_pages.append(
                        f"[Page {page_number}]\n{page_text}"
                    )

    except TextExtractionError:
        raise

    except Exception as error:
        raise TextExtractionError(
            f"Unable to open the PDF file: {error}"
        ) from error

    return "\n\n".join(extracted_pages)


def extract_scanned_pdf_text(path: Path) -> str:
    extracted_pages = []

    try:
        pdf_document = pdfium.PdfDocument(str(path))
    except Exception as error:
        raise TextExtractionError(
            f"Unable to open the scanned PDF: {error}"
        ) from error

    try:
        for page_index in range(len(pdf_document)):
            page = pdf_document[page_index]

            try:
                bitmap = page.render(
                    scale=PDF_RENDER_SCALE,
                )

                image = bitmap.to_pil()

                page_text = run_easyocr(image)

                if page_text.strip():
                    extracted_pages.append(
                        f"[Page {page_index + 1}]\n"
                        f"{page_text}"
                    )

            except Exception as error:
                raise TextExtractionError(
                    f"OCR failed on PDF page "
                    f"{page_index + 1}: {error}"
                ) from error

            finally:
                page.close()

    finally:
        pdf_document.close()

    if not extracted_pages:
        raise TextExtractionError(
            "No readable text was found in the scanned PDF"
        )

    return "\n\n".join(extracted_pages)


def extract_docx_text(path: Path) -> str:
    try:
        document = Document(str(path))
    except Exception as error:
        raise TextExtractionError(
            f"Unable to open the Word document: {error}"
        ) from error

    extracted_parts = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            extracted_parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            row_values = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_values.append(cell_text)

            if row_values:
                extracted_parts.append(
                    " | ".join(row_values)
                )

    for relationship in document.part.rels.values():
        target = getattr(
            relationship,
            "target_part",
            None,
        )

        content_type = getattr(
            target,
            "content_type",
            "",
        )

        if target and content_type.startswith("image/"):
            try:
                image = Image.open(
                    target.blob
                )

                image_text = run_easyocr(image)

                if image_text.strip():
                    extracted_parts.append(image_text)

            except Exception:
                # Embedded-image OCR is optional.
                # A failure here should not stop normal DOCX text extraction.
                continue

    return "\n".join(extracted_parts)


def extract_image_text(path: Path) -> str:
    try:
        with Image.open(path) as image:
            image_copy = image.convert("RGB")

        return run_easyocr(image_copy)

    except TextExtractionError:
        raise

    except Exception as error:
        raise TextExtractionError(
            f"Unable to process the image: {error}"
        ) from error


def run_easyocr(image: Image.Image) -> str:
    reader = get_ocr_reader()

    try:
        results = reader.readtext(
            image,
            detail=0,
            paragraph=False,
        )
    except Exception as error:
        raise TextExtractionError(
            f"EasyOCR could not read the image: {error}"
        ) from error

    readable_lines = []

    for result in results:
        cleaned_result = " ".join(
            str(result).split()
        )

        if cleaned_result:
            readable_lines.append(cleaned_result)

    return "\n".join(readable_lines)


def clean_extracted_text(text: str) -> str:
    cleaned_lines = []

    normalized_text = (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
    )

    for raw_line in normalized_text.split("\n"):
        stripped_line = raw_line.strip()

        if not stripped_line:
            continue

        if stripped_line.startswith("[Page ") and stripped_line.endswith("]"):
            cleaned_lines.append(stripped_line)
            continue

        cleaned_line = " ".join(
            stripped_line.split()
        )

        if cleaned_line:
            cleaned_lines.append(cleaned_line)

    return "\n".join(cleaned_lines).strip()